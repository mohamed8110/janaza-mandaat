import streamlit as st
from docx import Document
from datetime import date
import os

st.set_page_config(page_title="Mandaatgenerator - Janaza", layout="centered")
st.title("🖊️ Mandaat / Volmacht")

st.markdown("Vul de onderstaande gegevens in om een mandaatdocument te genereren.")

with st.form("mandaat_form"):
    st.header("📌 Gegevens overledene")
    naam_overledene = st.text_input("Naam en voornaam")
    rijksregnr_overledene = st.text_input("Rijksregisternummer")
    geboortedatum = st.text_input("Geboortedatum (bv. 14/03/1967)")
    geboorteplaats = st.text_input("Geboorteplaats")
    adres_overledene = st.text_input("Adres")
    nationaliteit = st.text_input("Nationaliteit")
    burgerlijke_staat = st.text_input("Burgerlijke staat")
    datum_overlijden = st.text_input("Datum van overlijden")
    plaats_overlijden = st.text_input("Plaats van overlijden")

    st.header("👤 Contactpersoon")
    naam_contact = st.text_input("Naam en voornaam (contactpersoon)")
    rijksregnr_contact = st.text_input("Rijksregisternummer (contactpersoon)")
    adres_contact = st.text_input("Adres (contactpersoon)")
    email = st.text_input("E-mailadres")
    telefoon = st.text_input("Telefoonnummer")
    bloedverwantschap = st.text_input("Bloedverwantschap met de overledene")

    st.header("☑️ Bevestigingen")
    check_correct = st.checkbox("✅ Ik bevestig dat de gegevens correct zijn")
    check_volmacht = st.checkbox("✅ Ik geef volmacht aan Janaza VZW")
    check_zorg = st.checkbox("✅ Ik geef toestemming voor de zorgen aan de overledene")

    plaats_mandaat = st.text_input("Plaats ondertekening", value="Antwerpen")
    datum_mandaat = st.date_input("Datum ondertekening", value=date.today())
    bestandsnaam = st.text_input("Bestandsnaam voor het document", value="mandaat")

    submitted = st.form_submit_button("📄 Genereer mandaatdocument")

if submitted:
    template_path = "template_mandaat.docx"
    if not os.path.exists(template_path):
        st.error("❗ Templatebestand ontbreekt.")
    else:
        doc = Document(template_path)
        vervangingen = {
            "<<NAAM_OVERLEDENE>>": naam_overledene,
            "<<RIJKSREG_NR_OVERLEDENE>>": rijksregnr_overledene,
            "<<GEBOORTEDATUM>>": geboortedatum,
            "<<GEBOORTEPLAATS>>": geboorteplaats,
            "<<ADRES_OVERLEDENE>>": adres_overledene,
            "<<NATIONALITEIT>>": nationaliteit,
            "<<BURGERLIJKE_STAAT>>": burgerlijke_staat,
            "<<DATUM_OVERLIJDEN>>": datum_overlijden,
            "<<PLAATS_OVERLIJDEN>>": plaats_overlijden,
            "<<NAAM_CONTACT>>": naam_contact,
            "<<RIJKSREG_NR_CONTACT>>": rijksregnr_contact,
            "<<ADRES_CONTACT>>": adres_contact,
            "<<EMAIL>>": email,
            "<<TELEFOON>>": telefoon,
            "<<BLOEDVERWANTSCHAP>>": bloedverwantschap,
            "<<CHECK_GEGEVENS_CORRECT>>": "☑" if check_correct else "☐",
            "<<CHECK_VOLMACHT>>": "☑" if check_volmacht else "☐",
            "<<CHECK_TOESTEMMING_ZORG>>": "☑" if check_zorg else "☐",
            "<<DATUM_MANDAAT>>": datum_mandaat.strftime("%d/%m/%Y"),
            "<<PLAATS_MANDAAT>>": plaats_mandaat
        }

        for para in doc.paragraphs:
            for run in para.runs:
                for key, val in vervangingen.items():
                    if key in run.text:
                        run.text = run.text.replace(key, val)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            for key, val in vervangingen.items():
                                if key in run.text:
                                    run.text = run.text.replace(key, val)

        output_path = f"{bestandsnaam}.docx"
        doc.save(output_path)

        with open(output_path, "rb") as f:
            st.success("✅ Mandaat gegenereerd!")
            st.download_button("📥 Download .docx", f, file_name=output_path, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")